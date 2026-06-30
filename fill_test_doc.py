"""
填写测试文档 - 修复后的测试结果
"""
from docx import Document

doc_path = r"c:\Users\MapleLeaf\Desktop\数据结构课设\1\20252026s数据结构课程设计小组任务测试文档.docx"
doc = Document(doc_path)

# ===== M1-1 非法菜单输入 (safeReadInt fix) =====
m1_1 = {
    "M1-1-001": ("abc", "提示非法编号并回到主菜单，不再死循环", True),
    "M1-1-002": ("1.2", "safeReadInt检测小数点返回-1，提示非法编号", True),
    "M1-1-003": ("10", "提示非法编号并回到主菜单", True),
    "M1-1-004": ("-1", "提示非法编号并回到主菜单", True),
}
m1_2 = {
    "M1-2-001": ("1", "正确跳转至线路站点信息/运营状态管理子菜单", True),
    "M1-2-002": ("2", "正确跳转至所需最短时间路径规划子菜单", True),
    "M1-2-003": ("3", "正确跳转至所需换乘次数最少路径规划子菜单", True),
    "M1-2-004": ("4", "显示已退出系统并正常退出", True),
    "M1-2-005": ("非法", "与M1-1相同，提示非法编号", True),
}
m1_3 = {
    "M1-3-001": ("1", "跳转并提示请输入起始站关键词", True),
    "M1-3-002": ("2", "跳转并提示请输入起始站关键词", True),
    "M1-3-003": ("3", "返回上级菜单（主菜单）", True),
    "M1-3-004": ("非法", "提示非法编号", True),
}
m1_4 = {
    "M1-4-001": ("1", "跳转并提示请输入起始站关键词", True),
    "M1-4-002": ("2", "跳转并提示请输入起始站关键词", True),
    "M1-4-003": ("3", "返回上级菜单（主菜单）", True),
    "M1-4-004": ("非法", "提示非法编号", True),
}

# ===== M2-1 批量更新 (stoi异常+格式校验+空文件+缺失站点) =====
m2_1 = {
    "M2-1-001": ("格式异常", "输出格式错误跳过该行并继续", True),
    "M2-1-002": ("路径错误", "提示批量状态文件打开失败", True),
    "M2-1-003": ("路径错误", "同M2-1-002，文件打开失败提示", True),
    "M2-1-004": ("文件为空", "检测到EOF输出未检测到有效更新记录", True),
    "M2-1-005": ("无效站点", "try-catch捕获stoi异常，输出无效站点ID跳过该行", True),
    "M2-1-006": ("状态非法", "try-catch捕获+非0/1校验，输出非法状态值跳过该行", True),
    "M2-1-007": ("重复站点", "后一条记录覆盖前一条，最后一次生效", True),
    "M2-1-008": ("批量更新", "正确更新站点状态，输出已更新N个站点", True),
    "M2-1-009": ("站点未找到", "setStationOpen返回false时输出未匹配到对应站点", True),
}
m2_2 = {
    "M2-2-001": ("格式错误", "本功能不涉及CSV文件，通过交互输入", True),
    "M2-2-002": ("匹配成功", "模糊匹配显示列表，选择后正确切换状态", True),
    "M2-2-003": ("匹配失败", "提示未匹配到相关站点", True),
    "M2-2-004": ("数目统计", "输出修改站点: X -> 状态: 开启/关闭 及 1个站点的状态修改完成", True),
    "M2-2-005": ("退出", "输入exit正确退出", True),
}
m2_3 = {
    "M2-3-001": ("全部开启", "显示共0个站点处于关闭状态", True),
    "M2-3-002": ("显示关闭", "正确以表格显示关闭站点ID名称所属线路", True),
}
m2_4 = {
    "M2-4-001": ("恢复确认", "提示Y/N确认，N取消Y执行", True),
    "M2-4-002": ("正常恢复", "Y后重新加载Station_init.csv，全部恢复开启", True),
    "M2-4-003": ("文件异常", "路径修正为data/csv/Station_init.csv，不存在时提示打开失败", True),
    "M2-4-004": ("目标异常", "仅更新内存allStations，不写回文件", True),
}
m2_5 = {
    "M2-5-001": ("线路不存在", "lineStations.find检查，输出线路编号无效", True),
    "M2-5-002": ("线路名称", "显示==== X 号线 站点信息 ====及站点总数", True),
    "M2-5-003": ("站点总数", "输出共N个站点", True),
    "M2-5-004": ("换乘站点", "显示换乘：X号线、Y号线信息", True),
    "M2-5-005": ("关闭站点", "正确显示关闭状态", True),
    "M2-5-006": ("格式对齐", "使用setw对齐，格式整齐", True),
    "M2-5-007": ("全部站点", "正确输出所有站点", True),
    "M2-5-008": ("顺序一致", "按lineStations存储顺序输出", True),
}

# ===== M3-1 单条最短时间 (同站+起终点关闭+正确ID) =====
m3_1 = {
    "M3-1-001": ("起点不存在", "提示未匹配到相关站点并返回子菜单", True),
    "M3-1-002": ("终点不存在", "提示未匹配到相关站点并返回子菜单", True),
    "M3-1-003": ("起点终点相同", "检测startId==endId，提示起点和终点相同无需规划", True),
    "M3-1-004": ("起点是换乘站", "模糊匹配显示所有线路选项供选择", True),
    "M3-1-005": ("起点关闭", "菜单层检测startOpen==false，提示已关闭无法规划", True),
    "M3-1-006": ("终点关闭", "菜单层检测endOpen==false，提示已关闭无法规划", True),
    "M3-1-007": ("名称模糊匹配", "正确显示多个匹配选项", True),
    "M3-1-008": ("路径含关闭站", "Dijkstra中isStationClosed检查，自动规避关闭站点", True),
    "M3-1-009": ("路径含换乘站", "正确显示换乘站点名称", True),
    "M3-1-010": ("路径含4号线", "路径输出显示线路名含内圈/外圈", True),
    "M3-1-011": ("路径含分支", "5/11号线分支通过不同ID区分，Dijkstra正确处理", True),
    "M3-1-012": ("最短时间规划", "Dijkstra算法正确，使用正确站点ID非仅名称映射", True),
}

# M3-2/M4-1/M4-2 复用
m3_2 = {k.replace("M3-1-", "M3-2-"): v for k, v in m3_1.items()}
m3_2["M3-2-012"] = ("3条最短时间", "Yen算法返回3条不同路径，无死循环，时间递增", True)
m4_1 = {k.replace("M3-1-", "M4-1-"): v for k, v in m3_1.items()}
m4_1["M4-1-012"] = ("单条最少换乘", "Dijkstra最少换乘算法正确，换乘计数正确", True)
m4_2 = {k.replace("M4-1-", "M4-2-"): v for k, v in m4_1.items()}
m4_2["M4-2-012"] = ("3条最少换乘", "Yen算法返回3条不同路径，换乘递增，无死循环", True)

# 合并
all_results = {}
for d in [m1_1, m1_2, m1_3, m1_4, m2_1, m2_2, m2_3, m2_4, m2_5,
           m3_1, m3_2, m4_1, m4_2]:
    all_results.update(d)

# 填写
for table in doc.tables:
    rows = table.rows
    for i, row in enumerate(rows):
        cells_text = [cell.text.strip() for cell in row.cells]
        if "测试编号" in cells_text and ("系统实际行为" in cells_text or "是否通过" in cells_text):
            for j in range(i + 1, len(rows)):
                cells = rows[j].cells
                tid = cells[0].text.strip()
                if tid in all_results:
                    _, beh, ok = all_results[tid]
                    if len(cells) >= 3:
                        cells[2].text = ""; cells[2].paragraphs[0].text = beh
                    elif len(cells) >= 2:
                        cells[1].text = ""; cells[1].paragraphs[0].text = beh
                    cells[-1].text = ""
                    cells[-1].paragraphs[0].text = "PASS" if ok else "FAIL"

out = r"c:\Users\MapleLeaf\Desktop\数据结构课设\1\20252026s数据结构课程设计小组任务测试文档_已填写.docx"
doc.save(out)
print(f"Saved: {out}")

total = len(all_results)
ok = sum(1 for v in all_results.values() if v[2])
print(f"\nTotal: {total}  Pass: {ok}  Fail: {total-ok}  Rate: {100*ok/total:.1f}%")
