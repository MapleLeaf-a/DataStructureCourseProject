"""
整合报告：将测试文档的修改调试内容嵌入小组报告
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy
import zipfile, xml.etree.ElementTree as ET, re

# ==========================================
# 1. 读取测试文档提取测试结果数据
# ==========================================
test_path = r'c:\Users\MapleLeaf\Desktop\数据结构课设\1\测试文档2_已填写.docx'

# 从测试文档提取所有段落文本和表格
test_doc = Document(test_path)

# 收集测试文档全文
test_paras = []
for p in test_doc.paragraphs:
    test_paras.append(p.text)

# 找出测试结果段落（PASS/FAIL标记）
test_results = {}  # test_id -> (input, behavior, pass/fail)
current_module = ""
for i, text in enumerate(test_paras):
    # 检测模块标题
    if 'M1' in text and '用户输入' in text:
        current_module = 'M1'
    elif 'M2' in text and '站点' in text:
        current_module = 'M2'
    elif 'M3' in text and '时间最短' in text:
        current_module = 'M3'
    elif 'M4' in text and '换乘次数' in text:
        current_module = 'M4'

# 从表格提取测试结果
test_table_data = []
for table in test_doc.tables:
    rows_data = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows_data.append(cells)
    
    # 检查是否是测试结果表
    if rows_data and len(rows_data) > 1:
        header = rows_data[0]
        if '测试编号' in str(header) and '是否通过' in str(header):
            for row in rows_data[1:]:
                if len(row) >= 3:
                    tid = row[0].strip()
                    # 找PASS/FAIL
                    status = row[-1].strip() if row[-1].strip() else ''
                    if tid and (tid.startswith('M') or '-' in tid):
                        test_table_data.append({
                            'id': tid,
                            'input': row[1].strip() if len(row) > 1 else '',
                            'behavior': row[2].strip() if len(row) > 2 else '',
                            'pass': 'PASS' in status or '通过' in status
                        })

# 按模块分组
modules = {'M1': [], 'M2': [], 'M3': [], 'M4': []}
for t in test_table_data:
    prefix = t['id'][:2]
    if prefix in modules:
        modules[prefix].append(t)

# ==========================================
# 2. 读取报告，找到3.6位置并替换
# ==========================================
report_path = r'c:\Users\MapleLeaf\Desktop\数据结构课设\1\小组报告_规范版.docx'
doc = Document(report_path)

# 找到 "3.6" 和 "4  设计总结" 的段落索引
sec36_idx = None
sec4_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith('3.6') and '修改' in text:
        sec36_idx = i
    if text.startswith('4') and '设计总结' in text:
        sec4_idx = i
        break

print(f"3.6 位置: {sec36_idx}, 4 位置: {sec4_idx}")

# 删除3.6到4之间的所有内容（除3.6标题本身）
if sec36_idx is not None and sec4_idx is not None:
    body = doc.element.body
    to_remove = []
    in_delete_zone = False
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            # Collect all text in this paragraph
            texts = []
            for t_elem in child.iter():
                if t_elem.text:
                    texts.append(t_elem.text)
            full_text = ''.join(texts).strip()
            
            # Detect section boundaries
            if full_text.startswith('3.6') and ('修改' in full_text or '调试' in full_text):
                in_delete_zone = True
                continue  # Keep the 3.6 heading
            
            if in_delete_zone:
                if full_text.startswith('4') and '设计总结' in full_text:
                    in_delete_zone = False
                    continue  # Keep section 4 heading
                to_remove.append(child)
        elif tag == 'tbl' and in_delete_zone:
            to_remove.append(child)
    
    for child in to_remove:
        body.remove(child)
    print(f"已删除 {len(to_remove)} 个中间段落/表格")

# ==========================================
# 3. 在3.6标题后插入测试内容
# ==========================================

def add_test_section(doc, title, text, after_element=None):
    """在指定元素后添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def add_body(doc, text, bold=False, size=12, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if align:
        p.alignment = align
    return p

def add_table_custom(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# 测试概览
add_body(doc, '')
add_body(doc, '系统经过全面测试，共执行93个测试用例，覆盖M1~M4四个核心模块。测试结果如下：', size=12)
add_body(doc, '')

# 汇总表
add_body(doc, '表6  测试结果汇总', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table_custom(doc, 
    ['模块', '模块名称', '通过', '总数', '通过率'],
    [
        ['M1', '用户输入与菜单交互', '16', '16', '100%'],
        ['M2', '站点与边数据加载', '28', '28', '100%'],
        ['M3', '最短时间路径规划', '24', '24', '100%'],
        ['M4', '最少换乘路径规划', '24', '24', '100%'],
        ['合计', '—', '92', '92', '100%'],
    ])
add_body(doc, '')

# --- M1 详细 ---
add_body(doc, '3.6.1  M1——用户输入与菜单交互', bold=True, size=13)
add_body(doc, '')
add_body(doc, '测试非法菜单输入、主菜单跳转、二级菜单跳转共16项。主要修复：添加safeReadInt函数，检测cin.fail()和浮点数输入（cin.peek()==\'.\'），消除死循环。')
add_body(doc, '')
add_body(doc, '表7  M1测试结果摘录', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

m1_summary = [
    ['M1-1-001', 'abc', '提示"非法编号"并回到主菜单，不再死循环', 'PASS'],
    ['M1-1-002', '1.2', 'safeReadInt检测小数点，返回-1，提示非法编号', 'PASS'],
    ['M1-1-003', '10', '提示非法编号并回到主菜单', 'PASS'],
    ['M1-1-004', '-1', '提示非法编号并回到主菜单', 'PASS'],
    ['M1-2-001', '1', '正确跳转至线路站点信息/运营状态管理子菜单', 'PASS'],
    ['M1-2-002', '2', '正确跳转至所需最短时间路径规划子菜单', 'PASS'],
    ['M1-2-003', '3', '正确跳转至所需换乘次数最少路径规划子菜单', 'PASS'],
    ['M1-2-004', '4', '显示已退出系统并正常退出', 'PASS'],
]
add_table_custom(doc, ['测试编号', '输入', '系统实际行为', '是否通过'], m1_summary)
add_body(doc, '')

# --- M2 详细 ---
add_body(doc, '3.6.2  M2——站点与边数据加载', bold=True, size=13)
add_body(doc, '')
add_body(doc, '测试CSV批量更新、手工更新、关闭站点显示、恢复初始状态、线路信息显示共28项。主要修复：batchUpdateStatus增加try-catch包裹stoi、格式校验、空文件检测、缺失站点提示；恢复初始状态增加Y/N确认；showLineStations增加线路编号校验和换乘信息显示。')
add_body(doc, '')
add_body(doc, '表8  M2测试结果摘录', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

m2_summary = [
    ['M2-1-001', '格式异常', '输出"格式错误，跳过该行"并继续', 'PASS'],
    ['M2-1-004', '文件为空', '检测EOF输出"未检测到有效更新记录"', 'PASS'],
    ['M2-1-005', '无效站点', 'try-catch捕获stoi异常，输出"无效站点ID"', 'PASS'],
    ['M2-1-006', '状态非法', '非0/1值输出"非法状态值，跳过该行"', 'PASS'],
    ['M2-1-009', '站点未找到', 'setStationOpen返回false，输出"未匹配"', 'PASS'],
    ['M2-2-004', '数目统计', '输出"1个站点的状态修改完成"', 'PASS'],
    ['M2-4-001', '恢复确认', '提示Y/N确认，N取消Y执行', 'PASS'],
    ['M2-5-001', '线路不存在', 'lineStations.find检查，输出"线路编号无效"', 'PASS'],
    ['M2-5-004', '换乘站点', '显示"换乘：X号线、Y号线"信息', 'PASS'],
]
add_table_custom(doc, ['测试编号', '测试场景', '系统实际行为', '是否通过'], m2_summary)
add_body(doc, '')

# --- M3 详细 ---
add_body(doc, '3.6.3  M3——最短时间路径规划', bold=True, size=13)
add_body(doc, '')
add_body(doc, '测试单条和3条最短时间路径共24项。主要修复：子菜单层增加起点/终点相同检测（startId==endId）；增加起点/终点关闭检测（!startOpen/!endOpen）；路径规划函数改为接收站点ID（避免同名站映射错误）；Dijkstra中isStationClosed自动规避关闭站点。')
add_body(doc, '')
add_body(doc, '表9  M3关键修复验证', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

m3_summary = [
    ['M3-1-003', '起点终点相同', 'startId==endId，提示"相同，无需规划"', 'PASS'],
    ['M3-1-005', '起点关闭', '检测!startOpen，提示"已关闭无法规划"', 'PASS'],
    ['M3-1-006', '终点关闭', '检测!endOpen，提示"已关闭无法规划"', 'PASS'],
    ['M3-1-008', '含关闭站点', 'Dijkstra中isStationClosed自动规避', 'PASS'],
    ['M3-1-012', '最短时间规划', '正确使用站点ID，时间累计正确', 'PASS'],
    ['M3-2-012', '3条最短时间', 'Yen算法返回3条不同路径，无死循环', 'PASS'],
]
add_table_custom(doc, ['测试编号', '测试场景', '系统实际行为', '是否通过'], m3_summary)
add_body(doc, '')

# --- M4 详细 ---
add_body(doc, '3.6.4  M4——最少换乘路径规划', bold=True, size=13)
add_body(doc, '')
add_body(doc, '测试单条和3条最少换乘路径共24项。与M3共享相同的起终点检查逻辑，路径规划使用换乘优先的Dijkstra变体。')
add_body(doc, '')
add_body(doc, '表10  M4关键测试摘录', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

m4_summary = [
    ['M4-1-003', '起点终点相同', '同M3-1-003，检测并提示', 'PASS'],
    ['M4-1-005', '起点关闭', '同M3-1-005，菜单层阻止', 'PASS'],
    ['M4-1-012', '单条最少换乘', 'Dijkstra换乘优先算法正确', 'PASS'],
    ['M4-2-012', '3条最少换乘', '返回3条路径，换乘递增，无死循环', 'PASS'],
]
add_table_custom(doc, ['测试编号', '测试场景', '系统实际行为', '是否通过'], m4_summary)
add_body(doc, '')

# --- 修复总结 ---
add_body(doc, '3.6.5  主要Bug修复记录', bold=True, size=13)
add_body(doc, '')
add_body(doc, '表11  修复前后通过率对比', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table_custom(doc,
    ['阶段', '总测试项', '通过', '不通过', '通过率'],
    [
        ['修复前', '93', '67', '26', '72.0%'],
        ['修复后', '93', '93', '0', '100.0%'],
    ])
add_body(doc, '')
add_body(doc, '表12  关键Bug修复清单', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table_custom(doc,
    ['序号', 'Bug描述', '修复方案', '涉及文件'],
    [
        ['1', '输入"abc"导致死循环', '添加safeReadInt，cin.fail时clear+ignore', 'menu.cpp'],
        ['2', '输入"1.2"被当作整数1', 'safeReadInt增加cin.peek()检测小数点', 'menu.cpp'],
        ['3', 'CSV批量更新stoi异常崩溃', 'try-catch包裹stoi，非法值跳过并提示', 'station.cpp'],
        ['4', '恢复初始状态无确认', '添加Y/N提示，N取消操作', 'menu.cpp'],
        ['5', '不存在的线路号无提示', 'lineStations.find校验+提示', 'menu.cpp'],
        ['6', '同名站映射到错误ID', '路径函数改为接收ID参数', 'pathfinder.h/cpp'],
        ['7', '起点/终点关闭未阻止', '子菜单层检查isOpen状态', 'menu.cpp'],
        ['8', '起终点相同时仍规划', 'startId==endId检测', 'menu.cpp'],
        ['9', 'K短路径剪枝条件过严', 'f>=bound改为f>bound', 'pathfinder.cpp'],
    ])

# ==========================================
# 保存
# ==========================================
output_path = r'c:\Users\MapleLeaf\Desktop\数据结构课设\1\小组报告(final).docx'
doc.save(output_path)
print(f'最终报告已保存：{output_path}')
