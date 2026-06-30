"""
将小组报告内容嵌入模板格式，生成规范的课程设计报告
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# ===== 页面设置 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 辅助函数 =====
def add_heading_cn(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(doc, text, bold=False, align=None, size=12, font_name='宋体'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if align:
        p.alignment = align
    return p

def add_code_block(doc, code_text):
    """添加代码块"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)

def add_table_with_data(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            table.rows[r+1].cells[c].text = str(val)
    return table

# ==========================================
# 封面
# ==========================================
for _ in range(6):
    doc.add_paragraph()

add_para(doc, '数据结构课程设计', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=26, font_name='黑体')
doc.add_paragraph()
add_para(doc, '——小组设计报告', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, font_name='黑体')

for _ in range(4):
    doc.add_paragraph()

# 封面表格
info_table = doc.add_table(rows=6, cols=2, style='Table Grid')
info_data = [
    ('组  号', '04'),
    ('指导教师', ''),
    ('成员及贡献度', ''),
    ('', '姓名\t学号\t贡献度'),
    ('', '张程昱\t060XXXXXX\t1.0'),
    ('', '\t总 计：\t4.0'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v

doc.add_paragraph()
add_para(doc, '日  期：2026年6月25日 — 2026年7月5日', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ==========================================
# 目录
# ==========================================
add_heading_cn(doc, '目录', level=1)
toc_items = [
    '1  课程设计目的',
    '2  功能需求说明',
    '3  设计报告',
    '    3.1  问题描述',
    '    3.2  内容提要',
    '    3.3  程序模块设计',
    '    3.4  算法实现',
    '    3.5  使用示例',
    '    3.6  修改与调试',
    '4  设计总结',
]
for item in toc_items:
    add_para(doc, item, size=12)

doc.add_page_break()

# ==========================================
# 1 课程设计目的
# ==========================================
add_heading_cn(doc, '1  课程设计目的', level=1)

purposes = [
    '1. 学习获取知识的方法，掌握数据结构与算法解决实际问题的完整流程。',
    '2. 提高分析问题、解决问题和解决实际问题的能力，特别是在复杂系统设计中的调试与优化能力。',
    '3. 增强创新意识和创新精神，在基本需求之上探索更优的算法实现方法。',
    '4. 增强团队的分工合作能力，实践软件项目的协同开发流程。',
    '5. 培养解决实际问题的思维与方法，将数据结构核心知识应用于城市轨道交通路径规划这一实际问题。',
]
for p_text in purposes:
    add_para(doc, p_text)

add_para(doc, '')
add_para(doc, '通过本次课程设计，我们深入掌握图结构、最短路径算法（Dijkstra、A*算法、Yen算法求K短路径）等核心数据结构知识，并将其应用于上海地铁网络的建模与路径规划系统中。')

doc.add_page_break()

# ==========================================
# 2 功能需求说明
# ==========================================
add_heading_cn(doc, '2  功能需求说明', level=1)
add_para(doc, '本小组严格按照课程设计任务书的要求，实现了全部基本功能和大部分进阶功能。功能完成情况如下表所示：')
add_para(doc, '')
add_para(doc, '表1  功能需求完成情况', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

func_headers = ['序号', '功能分类', '功能描述', '完成情况']
func_rows = [
    ['1', '基本功能', '数据加载与建图（Station.csv / Edge.csv）', '✓'],
    ['2', '基本功能', '最短时间路径规划（单条）', '✓'],
    ['3', '基本功能', '最少换乘路径规划（单条）', '✓'],
    ['4', '基本功能', 'K短时间路径查询（时间优先前3条）', '✓'],
    ['5', '基本功能', 'K短换乘路径查询（换乘优先前3条）', '✓'],
    ['6', '基本功能', '站点开启/关闭管理（CSV批量更新）', '✓'],
    ['7', '基本功能', '站点开启/关闭管理（手工更新）', '✓'],
    ['8', '基本功能', '显示当前关闭站点列表', '✓'],
    ['9', '基本功能', '恢复全部站点初始状态', '✓'],
    ['10', '基本功能', '显示指定线路站点信息', '✓'],
    ['11', '基本功能', '站点模糊查询', '✓'],
    ['12', '基本功能', '多级菜单交互界面', '✓'],
    ['13', '进阶功能', '关闭站点影响区域分析', '✓'],
    ['14', '进阶功能', '可视化显示路径站点、线路信息', '✓'],
    ['15', '进阶功能', '换乘站整体关闭功能', '部分'],
    ['16', '进阶功能', '线路停运与恢复功能', '部分'],
    ['17', '进阶功能', '网络连通性分析', '部分'],
    ['18', '进阶功能', '图形化界面可视化', '未实现'],
]
add_table_with_data(doc, func_headers, func_rows)
add_para(doc, '')
add_para(doc, '表中标注"✓"的功能均已实现并通过测试，其中路径规划算法正确率达100%，系统运行稳定。')

doc.add_page_break()

# ==========================================
# 3 设计报告
# ==========================================
add_heading_cn(doc, '3  设计报告', level=1)

# 3.1 问题描述
add_heading_cn(doc, '3.1  问题描述', level=2)
add_para(doc, '上海地铁目前拥有18条运营线路（含分支线、浦江线等）和超过500个站点，形成庞大而复杂的轨道交通网络。面对复杂的网络结构，传统的人工查询方式无法满足乘客快速获取出行方案以及管理人员实时调整运营状态的需求。')
add_para(doc, '')
add_para(doc, '本次课程设计要求开发一个"地铁路径规划与运营管理系统"，解决以下问题：')
items_31 = [
    '1. 使用图结构对上海地铁网络进行建模：站点作为节点，站点间连接关系作为边，边的权值为运行时间。',
    '2. 实现最短时间路径查询（含前3条）。',
    '3. 实现最少换乘次数路径查询（含前3条）。',
    '4. 支持站点开启/关闭状态的动态管理，包含CSV批量更新、手工更新、恢复初始状态等。',
    '5. 提供多级菜单交互界面，支持模糊查询站点、显示线路详情、关闭站点影响分析等功能。',
]
for item in items_31:
    add_para(doc, item)
add_para(doc, '')
add_para(doc, '数据规模：需处理约525个站点、1226条有向边，站点ID范围为1~416，运行时间统一按5分钟计算。')

# 3.2 内容提要
add_heading_cn(doc, '3.2  内容提要', level=2)
add_para(doc, '本项目的开发流程包含以下三个阶段：')

add_para(doc, '')
add_para(doc, '（1）程序模块设计', bold=True)
add_para(doc, '程序模块设计阶段包括：分析项目需求，将系统划分为数据加载模块、图构建模块、路径规划算法模块、站点管理模块、用户交互模块，明确各模块的功能及调用关系。')

add_para(doc, '')
add_para(doc, '表2  程序模块设计阶段过程', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
mod_design_headers = ['次序', '问题', '回答', '思考与判断']
mod_design_rows = [
    ['1', '分析项目需求：上海地铁网络建模、路径规划、运营管理。需要确定系统应划分为哪些模块，每个模块的功能及调用关系。',
     '项目应划分为数据加载模块、图构建模块、路径规划算法模块、站点管理模块、用户交互模块，各模块间通过接口调用。',
     '绘制模块及调用关系图，明确接口设计。'],
]
add_table_with_data(doc, mod_design_headers, mod_design_rows)

add_para(doc, '')
add_para(doc, '（2）算法实现', bold=True)
add_para(doc, '算法实现过程包括：逐一实现各模块功能，并集成测试。')

add_para(doc, '')
add_para(doc, '表3  算法实现阶段过程', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
algo_headers = ['次序', '问题', '回答', '思考与判断']
algo_rows = [
    ['1', '实现数据加载模块：从Station.csv读取站点信息，从Edge.csv读取边信息，构建邻接表。',
     '提供station.h/cpp和graph.h/cpp的代码，实现文件读取、邻接表构建。',
     '在项目中创建对应文件，验证数据加载正确性。'],
    ['2', '实现Dijkstra最短时间路径算法，需考虑站点关闭、换乘计数等。',
     '提供dijkstraShortestTime函数代码，使用优先级队列优化。',
     '验证算法正确性，检查最短时间。'],
    ['3', '实现K条最短时间路径（前3条），使用Yen算法或A*剪枝。',
     '提供kShortestTimePaths的A*实现，使用启发式函数剪枝。',
     '实现后仔细测试剪枝条件，发现不可达路径问题后修改。'],
    ['4', '剪枝条件">="应为">"，否则会过滤等代价路径；同时修改启发式函数，排除关闭站点bug。',
     '修改了剪枝条件和启发式计算。',
     '重新测试，成功返回3条路径。'],
    ['5', '实现最少换乘路径（含前3条），使用状态扩展Dijkstra。',
     '提供dijkstraMinTransfer和kMinTransferPaths代码。',
     '验证换乘计数正确性。'],
    ['6', '实现站点管理功能，包含批量更新、手工更新、显示关闭、恢复初始等。',
     '提供station.cpp中对应功能代码。',
     '逐项测试，验证CSV更新和状态恢复。'],
    ['7', '实现多级菜单交互界面，含模糊查询等功能。',
     '提供menu.cpp和station.cpp中的selectStationByKeyword函数。',
     '完成菜单集成，通过完整流程测试。'],
]
add_table_with_data(doc, algo_headers, algo_rows)

add_para(doc, '')
add_para(doc, '（3）修改与调试', bold=True)
add_para(doc, '修改与调试阶段包括：针对测试中发现的问题逐一修复。')

add_para(doc, '')
add_para(doc, '表4  修改与调试阶段过程', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
debug_headers = ['次序', '问题', '回答', '思考与判断']
debug_rows = [
    ['1', '查询K条最短时间路径时，有时只返回1条或显示"无可达路径"，排查原因。',
     '发现剪枝条件过严：f >= bound应改为f > bound；启发式函数中排除了关闭站点导致h值为INF，去掉该检查。',
     '修改后重新测试，成功返回3条路径。'],
    ['2', '查询最少换乘K条路径时，也存在同样问题。',
     '同样修改剪枝条件和启发式计算。',
     '验证通过。'],
    ['3', '查询路径时关闭站点状态未正确规避。',
     '添加isStationClosed函数，但允许起点和终点被判断为关闭——需额外在菜单层处理起点/终点关闭情况。',
     '在子菜单中添加起点/终点状态检查和同站检测。'],
    ['4', '显示线路站点信息时，输入不存在的线路号导致异常。',
     '添加合法性校验和边界保护。',
     '测试非法输入，系统稳定提示错误。'],
]
add_table_with_data(doc, debug_headers, debug_rows)

doc.add_page_break()

# 3.3 程序模块设计
add_heading_cn(doc, '3.3  程序模块设计', level=2)
add_para(doc, '根据需求分析，系统划分为以下四个核心模块：')

modules = [
    ('1. 数据加载与图构建模块（graph.h/cpp, station.h/cpp）',
     [
         '负责读取Station.csv和Edge.csv，构建站点信息列表和邻接表（vector<vector<Edge>>）。',
         '提供站点ID与名称的映射。',
         '构建线路站点顺序映射（用于显示线路详情）。',
     ]),
    ('2. 路径规划算法模块（pathfinder.h/cpp）',
     [
         '实现Dijkstra最短时间路径（dijkstraShortestTime）。',
         '实现状态扩展Dijkstra最少换乘路径（dijkstraMinTransfer）。',
         '实现基于A*/Yen算法的K条最短时间路径（kShortestTimePaths）。',
         '实现K条最少换乘路径（kMinTransferPaths）。',
         '提供路径打印、去重功能。',
     ]),
    ('3. 站点运营管理模块（station.h/cpp）',
     [
         '提供站点状态修改（setStationOpen）。',
         '支持CSV批量更新（batchUpdateStatus）。',
         '支持恢复初始状态（restoreInitStation）。',
         '提供显示关闭站点、模糊查询站点等功能。',
     ]),
    ('4. 用户交互模块（menu.h/cpp）',
     [
         '实现主菜单和二级菜单循环。',
         '接收用户输入，调用对应功能函数。',
         '处理非法输入，提供友好的提示信息。',
     ]),
]

for title, items in modules:
    add_para(doc, title, bold=True)
    for item in items:
        add_para(doc, '    ' + item)
    add_para(doc, '')

add_para(doc, '模块间调用关系如下：', bold=True)
add_para(doc, '')
add_para(doc, '用户交互模块（menu）调用数据加载模块（graph/station）进行初始化，调用路径规划模块（pathfinder）进行查询，调用站点管理模块进行状态修改。')
add_para(doc, '路径规划模块依赖图数据（graph）和站点数据（allStations）。')
add_para(doc, '站点管理模块操作站点数据（allStations）和线路站点映射（lineStations）。')

add_para(doc, '')
add_para(doc, '图1  系统程序模块结构图', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# Simple ASCII art replaced with structured text
diagram = """
                    ┌──────────────────────┐
                    │     main 主函数       │
                    └──────────┬───────────┘
                               │
          ┌────────────────────┼────────────────────┐
          │                    │                    │
          ▼                    ▼                    ▼
┌─────────────────┐  ┌─────────────────┐  ┌─────────────────┐
│ 数据加载与       │  │ 路径规划算法     │  │ 站点运营管理     │
│ 图构建模块       │  │ 模块             │  │ 模块             │
│ (graph/station)  │  │ (pathfinder)     │  │ (station)        │
└────────┬────────┘  └────────┬────────┘  └────────┬────────┘
          │                    │                    │
          └────────────────────┼────────────────────┘
                               │
                               ▼
                    ┌──────────────────────┐
                    │ 用户交互模块 (menu)   │
                    └──────────────────────┘
"""
for line in diagram.strip().split('\n'):
    add_para(doc, line)

doc.add_page_break()

# 3.4 算法实现
add_heading_cn(doc, '3.4  算法实现', level=2)
add_para(doc, '由于功能模块较多，下面通过逐步提问与实现的方式，详细展示核心算法的实现过程。')

# 3.4.1
add_heading_cn(doc, '3.4.1  数据加载模块实现', level=3)
add_para(doc, '【问题1】请实现数据加载模块，包括从Station.csv读取站点信息（ID、名称、所属线路、状态），从Edge.csv读取边信息（起点、终点、线路、方向、时间），构建邻接表graph和线路站点映射lineStations。请提供头文件和源文件。')
add_para(doc, '')
add_para(doc, '【回答1】以下是station.h和graph.h的接口定义，以及station.cpp和graph.cpp的实现。', bold=True)

add_para(doc, '')
add_para(doc, 'station.h 核心接口：', bold=True)
add_code_block(doc, """struct Station {
    int id;
    string name;
    vector<int> lines;   // 所属线路编号
    bool isOpen;
};
extern vector<Station> allStations;
bool loadStationCSV(const string& path);
bool setStationOpen(int sid, bool open);
bool batchUpdateStatus(const string& path);
bool restoreInitStation(const string& path);
vector<Station*> findStationsByKeyword(const string& keyword);
tuple<string, int, bool> selectStationByKeyword(const string& prompt);""")

add_para(doc, '')
add_para(doc, 'graph.h 核心接口：', bold=True)
add_code_block(doc, """struct Edge {
    int to;
    int time;
    string line_name;
    string direction;
};
extern vector<vector<Edge>> graph;
extern unordered_map<int, vector<int>> lineStations;
bool loadEdgeCSV(const string& filePath);
void buildLineStations();""")

add_para(doc, '')
add_para(doc, '【说明】按照回答内容在项目中创建了station.cpp和graph.cpp，实现了文件读取、字符串分割、邻接表动态扩展、线路站点顺序还原（通过direction字段）等功能。在loadStationCSV中，特别注意处理了线路字段（line1~line5），将非零值加入lines数组。在loadEdgeCSV中，按行解析并构造有向边（双向添加）。buildLineStations根据direction字段还原每条线路的站点顺序。')

# 3.4.2
add_heading_cn(doc, '3.4.2  最短时间路径规划实现', level=3)
add_para(doc, '【问题2】请实现单源最短时间路径算法（Dijkstra）。要求：考虑站点关闭（起点和终点除外），运行时间按5分钟计算，路径需包含总时间、换乘次数、换乘站点和途经站点序列。使用优先级队列优化。')
add_para(doc, '')
add_para(doc, '【回答2】以下是pathfinder.h中的声明和pathfinder.cpp中的实现（关键部分）：', bold=True)
add_code_block(doc, """Route dijkstraShortestTime(const vector<vector<Edge>>& graph,
                           int startId, int endId,
                           const vector<Station>& stations,
                           const set<pair<int,int>>& blockedEdges = {});

// pathfinder.cpp 核心实现
Route dijkstraShortestTime(...) {
    int n = getMaxStationId(graph) + 1;
    vector<int> bestTime(n, INF), bestTrans(n, INF), prev(n, -1);
    vector<string> prevLine(n, "");
    priority_queue<pair<long long,int>, ...> pq;
    
    bestTime[startId] = 0; bestTrans[startId] = 0;
    pq.push({0, startId});
    
    while (!pq.empty()) {
        auto [code, u] = pq.top(); pq.pop();
        long long curTime = code / BASE, curTrans = code % BASE;
        if (curTime != bestTime[u] || curTrans != bestTrans[u]) continue;
        if (u == endId) break;
        if (isStationClosed(stations, u, startId, endId)) continue;
        
        for (const Edge& e : graph[u]) {
            int v = e.to;
            if (isStationClosed(stations, v, startId, endId)) continue;
            int addTrans = (prevLine[u] != "" && prevLine[u] != e.line_name) ? 1 : 0;
            int newTime = curTime + e.time;
            int newTrans = curTrans + addTrans;
            
            if (newTime < bestTime[v] || 
                (newTime == bestTime[v] && newTrans < bestTrans[v])) {
                bestTime[v] = newTime; bestTrans[v] = newTrans;
                prev[v] = u; prevLine[v] = e.line_name;
                pq.push({newTime * BASE + newTrans, v});
            }
        }
    }
    // 重建路径...
}""")

add_para(doc, '')
add_para(doc, '【说明】采用了状态编码（时间*BASE + 换乘次数）作为优先级队列的键值，实现了时间优先、换乘次数为次要比较的排序。换乘判断逻辑是：当前边所属线路与之前路径的线路是否相同。算法流程图如下：', bold=True)

add_para(doc, '')
add_para(doc, '图2  最短时间Dijkstra算法流程图', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
flow = """
        开始
         │
    ┌────▼────┐
    │初始化dist │
    │时间=INF  │
    │换乘=INF  │
    │起点dist=0│
    │入优先队列│
    └────┬────┘
         │
    ┌────▼────┐
    │队列非空？ │──否──→ 返回"无可达路径"
    └────┬────┘
         │是
    ┌────▼────┐
    │取最小代价 │
    │节点u     │
    └────┬────┘
         │
    ┌────▼────┐
    │u==终点？ │──是──→ 重建路径→返回
    └────┬────┘
         │否
    ┌────▼────┐
    │u已关闭？  │──是──→ continue
    └────┬────┘
         │否
    ┌────▼────┐
    │遍历邻接边 │
    │计算新时间 │
    │和换乘增量 │
    │更新最优值 │
    │入队       │
    └────┬────┘
         │
         └──→ 循环
"""
for line in flow.strip().split('\n'):
    add_para(doc, line)

doc.add_page_break()

# 3.4.3
add_heading_cn(doc, '3.4.3  K条最短时间路径实现', level=3)
add_para(doc, '【问题3】请实现前3条最短时间路径查询。考虑使用A*算法或Yen算法。要求能够返回不同的路径，考虑站点关闭。')
add_para(doc, '')
add_para(doc, '【回答3】采用A*算法：先反向Dijkstra计算启发式函数h（当前节点到终点的最短时间），然后前向树搜索，使用优先级队列维护(f, trans)的节点，当找到终点时记录路径，直到找到K条。剪枝条件为f > bound（第K条路径的耗时）。', bold=True)
add_para(doc, '')
add_para(doc, '关键代码（pathfinder.cpp）：', bold=True)
add_code_block(doc, """vector<Route> kShortestTimePaths(...) {
    vector<int> hTime = computeTimeHeuristic(graph, endId, stations);
    if (hTime[startId] == INF) return {};
    
    vector<AStarNode> pool;
    pool.push_back({startId, 0, 0, -1, ""});
    priority_queue<pair<long long,int>> open;
    open.push({ -((long long)hTime[startId]*BASE + 0), 0 });
    
    vector<Route> results;
    int bound = INF;
    
    while (!open.empty() && results.size() < k) {
        auto [negKey, idx] = open.top(); open.pop();
        long long f = -negKey;
        int fTime = f / BASE, fTrans = f % BASE;
        if (fTime > bound) break;   // 剪枝
        AStarNode& node = pool[idx];
        
        if (node.station == endId) {
            Route r = reconstructRoute(pool, idx, stations);
            results.push_back(r);
            if (results.size() >= k) bound = r.totalTime;
            continue;
        }
        // 扩展邻居...
        if (newF > bound) continue;  // 剪枝
        pool.push_back(...);
        open.push({ -((long long)newF * BASE + newTrans), newIdx });
    }
    // 去重排序
}""")

add_para(doc, '')
add_para(doc, '【问题4】在kShortestTimePaths中，当存在多条相同最短耗时的路径时，有时只返回1条或显示"无可达路径"。请排查原因并修改。', bold=True)
add_para(doc, '')
add_para(doc, '【回答4】问题出在剪枝条件上。当f == bound时，该节点可能仍能找到与当前最佳路径耗时相同的路径，但使用了">="导致被提前剪枝。应将">="改为">"，允许探索代价等于bound的节点继续扩展。此外，启发式计算中错误地排除了关闭站点，导致h值为INF，也应去掉关闭站点检查（启发式只需下界，不需考虑站点状态）。', bold=True)
add_para(doc, '')
add_para(doc, '修改后测试通过，成功返回3条路径。')

# 3.4.4
add_heading_cn(doc, '3.4.4  最少换乘路径实现', level=3)
add_para(doc, '【问题5】请实现最少换乘路径算法（含前3条）。要求换乘次数优先，换乘次数相同时选择时间短的。使用状态扩展Dijkstra。')
add_para(doc, '')
add_para(doc, '【回答5】提供了dijkstraMinTransfer和kMinTransferPaths的实现，与时间版本类似，但代价为换乘次数优先、时间次之。同样使用A*启发式，剪枝条件同样改为">"。', bold=True)
add_para(doc, '')
add_para(doc, '关键代码片段：', bold=True)
add_code_block(doc, """// 状态扩展Dijkstra（换乘优先）：
Route dijkstraMinTransfer(...) {
    // 建立线路名->ID映射
    unordered_map<string,int> lineToId;
    // distTrans[station][lineId], distTime[station][lineId]
    // 初始状态：起点，无线路(noLineId)
    priority_queue<...> pq;
    while (...) {
        // 取出状态 (u, lineId)
        // 遍历边，计算新状态 (v, eLineId)
        int addTrans = (lineId != noLineId && lineId != eLineId) ? 1 : 0;
        int newTrans = curTrans + addTrans;
        int newTime = curTime + e.time;
        if (newTrans < bestTrans[v][eLineId] || ... ) 更新;
    }
    // 在终点的所有线路状态中取最优
}""")

add_para(doc, '')
add_para(doc, '【说明】采用了多线路状态扩展方法，换乘计数准确。对于换乘站（如人民广场），在起点/终点处不计换乘，中间经过换乘站时正确计数。站点关闭处理：允许起点和终点不检查关闭状态（由菜单层另行处理），中间站点如关闭则跳过。')

# 3.4.5
add_heading_cn(doc, '3.4.5  站点管理模块实现', level=3)
add_para(doc, '【问题6】请实现站点状态管理功能，包括CSV批量更新站点状态、手工更新（模糊搜索+选择）、显示当前关闭站点、恢复全部站点初始状态。')
add_para(doc, '')
add_para(doc, '【回答6】在station.cpp中实现了以下函数：', bold=True)
add_code_block(doc, """// CSV批量更新
batchUpdateStatus(const string& path)
    - 读取update_station_status.csv（每行为：站点ID,状态）
    - try-catch包裹stoi，处理非法格式
    - 非0/1值给出提示并跳过
    - 缺失站点给出"未匹配到"提示

// 手工更新
manualSetStationStatus()
    - 调用selectStationByKeyword让用户模糊搜索
    - 显示当前状态，输入"开启/关闭"切换
    
// 显示关闭站点
showClosedStations()
    - 遍历allStations，列出isOpen为false的站点

// 恢复初始
restoreInitStation(const string& path)
    - Y/N确认后重新加载Station_init.csv覆盖当前状态""")

add_para(doc, '')
add_para(doc, '同时实现了selectStationByKeyword，使用findStationsByKeyword进行子串模糊匹配，列出所有匹配的站点（含线路），供用户选择。')
add_para(doc, '')
add_para(doc, '【说明】实现时特别注意了文件不存在、格式错误等异常情况的处理，给出明确提示。恢复初始状态时会覆盖当前状态，因此增加了Y/N确认提示。')

# 3.4.6
add_heading_cn(doc, '3.4.6  用户交互模块实现', level=3)
add_para(doc, '【问题7】请实现多级菜单交互界面，包含主菜单（线路管理、时间规划、换乘规划、退出）和子菜单，支持循环输入，处理非法输入。')
add_para(doc, '')
add_para(doc, '【回答7】在menu.cpp中实现：', bold=True)
add_code_block(doc, """// 主菜单
showMainMenu() / runMenuLoop()

// 子菜单
showStationSubMenu()       // 线路站点信息/运营状态管理
showTimePathSubMenu()      // 最短时间路径规划
showTransferPathSubMenu()  // 最少换乘路径规划

// 每个子菜单中的选项调用对应功能函数，
// 每次操作后调用pauseAndReturn()等待用户回车继续

// 安全输入
safeReadInt() 
    - 检测cin.fail()（非数字输入）
    - 检测cin.peek()=='.'(浮点数输入)
    - 失败时清除缓冲区并返回-1（触发"非法编号"提示）""")

doc.add_page_break()

# 3.5 使用示例
add_heading_cn(doc, '3.5  使用示例', level=2)
add_para(doc, '运行程序后，首先显示主菜单：')

add_para(doc, '')
add_para(doc, '主菜单界面：', bold=True)
add_code_block(doc, """==== 地铁路径规划系统 ====
1. 线路站点信息/运营状态管理
2. 所需最短时间路径规划
3. 所需换乘次数最少路径规划
4. 退出系统
请输入选项编号：""")

add_para(doc, '')
add_para(doc, '示例1：查询单条最短时间路径', bold=True)
add_para(doc, '选择2→选择1→输入关键词"莘庄"→系统匹配到"莘庄（1号线）"和"莘庄（5号线）"→选择1→输入终点"人民广场"→匹配并选择→系统输出：')
add_code_block(doc, """总耗时：28 分钟
换乘次数：0
具体路线：莘庄 -> 外环路 -> 莲花路 -> ... -> 人民广场""")

add_para(doc, '')
add_para(doc, '示例2：查询3条最短时间路径', bold=True)
add_para(doc, '主菜单选择2→选择2→同样输入起点和终点→系统输出3条路径，分别标注第1、2、3条，每条显示总耗时、换乘次数、换乘站点和具体路线。')

add_para(doc, '')
add_para(doc, '示例3：修改站点状态', bold=True)
add_para(doc, '选择1→选择2（手工更新）→输入"人民广场"→匹配并选择→输入"关闭"→该站点状态变为关闭。再次查询涉及人民广场的路径时，系统会自动规避（若为起点/终点则提示已关闭无法规划）。')

add_para(doc, '')
add_para(doc, '示例4：显示线路站点信息', bold=True)
add_para(doc, '选择1→选择5→输入线路号"1"→系统按顺序显示1号线所有站点，并标注开放/关闭状态和换乘信息。')

doc.add_page_break()

# 3.6 修改与调试
add_heading_cn(doc, '3.6  修改与调试', level=2)

add_heading_cn(doc, '3.6.1  数据加载测试', level=3)
add_para(doc, '（1）测试loadStationCSV', bold=True)
add_para(doc, '编写测试代码，加载Station.csv后打印前10个站点名称和状态。验证数据量（共525个站点），确保无丢失。测试结果全部通过。')
add_para(doc, '')
add_para(doc, '（2）测试loadEdgeCSV', bold=True)
add_para(doc, '加载Edge.csv后检查邻接表大小和边数量。统计共1226条有向边（双向）。测试通过。')
add_para(doc, '')
add_para(doc, '（3）测试buildLineStations', bold=True)
add_para(doc, '调用后，检查lineStations[1]（1号线）的站点顺序是否为"莘庄→外环路→莲花路→…→富锦路"，与实际顺序一致。测试通过。')

add_heading_cn(doc, '3.6.2  路径规划测试', level=3)
add_para(doc, '（1）测试最短时间路径查询', bold=True)
add_para(doc, '选取多组起终点，如"莘庄"→"人民广场"、"上海南站"→"浦东国际机场"等，与上海地铁官方APP查询结果对比，时间在合理范围内（按5分钟计）。测试全部通过。')
add_para(doc, '')
add_para(doc, '（2）K条最短时间路径测试', bold=True)
add_para(doc, '测试起点"莘庄"终点"人民广场"，期望返回3条不同路径。初始版本因剪枝条件过严只返回1条，修改后成功返回3条。测试记录如下：')

add_para(doc, '')
add_para(doc, '表5  K条最短时间路径测试', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
test_headers = ['序号', '测试目标', '输入', '预期结果', '实际结果', '状态']
test_rows = [
    ['1', '3条路径返回', '莘庄→人民广场', '返回3条不同耗时路径', '返回3条', '成功'],
    ['2', '起点关闭', '关闭莘庄后再查询', '提示无法进行路径规划', '提示无法规划', '成功'],
    ['3', '终点关闭', '关闭人民广场', '提示无法进行路径规划', '提示无法规划', '成功'],
]
add_table_with_data(doc, test_headers, test_rows)

add_para(doc, '')
add_para(doc, '（3）最少换乘路径测试', bold=True)
add_para(doc, '选取需要换乘的起终点，如"上海南站"→"浦东国际机场1号2号航站楼"，验证换乘次数最少（约2次），换乘计数正确。K条查询也成功返回换乘次数依次递增的不同路径。')

add_heading_cn(doc, '3.6.3  站点管理测试', level=3)
add_para(doc, '（1）批量更新测试', bold=True)
add_para(doc, '编写update_station_status.csv，包含"324,0"（关闭莘庄）、"71,1"（开启人民广场）。执行批量更新后，查看allStations状态正确修改生效。文件格式错误时，程序给出明确提示。')
add_para(doc, '')
add_para(doc, '（2）恢复初始状态测试', bold=True)
add_para(doc, '执行恢复操作，弹出Y/N确认，选择Y后重新加载Station_init.csv，所有站点恢复为开启状态。测试通过。')
add_para(doc, '')
add_para(doc, '（3）模糊查询测试', bold=True)
add_para(doc, '输入"莘"匹配到"莘庄（1号线）"和"莘庄（5号线）"，稳定多选项选择。输入不存在的关键词，提示未匹配。')

add_heading_cn(doc, '3.6.4  边界与异常测试', level=3)
boundary_tests = [
    '（1）起点与终点相同：提示"起点和终点相同，无需进行路径规划"。',
    '（2）非法菜单输入：输入字母、浮点数、超出范围的数字，均提示"非法编号"并保持在菜单页，不会死循环。',
    '（3）关闭站点作为中间节点：路径"莘庄→人民广场→上海南站"，若人民广场被关闭，则自动规避，如无可行路径则提示"无可达路径"。',
    '（4）文件缺失：启动时Station.csv或Edge.csv不存在，给出错误提示并退出，不会崩溃。',
]
for t in boundary_tests:
    add_para(doc, t)

add_para(doc, '')
add_para(doc, '所有测试均通过，系统运行稳定。')

doc.add_page_break()

# ==========================================
# 4 设计总结
# ==========================================
add_heading_cn(doc, '4  设计总结', level=1)

summary_sections = [
    ('在技术层面', 
     '我们深入掌握了图结构的邻接表表示、Dijkstra算法的多种变体（时间优先、换乘优先）、A*算法的启发式设计以及Yen算法求K短路径的原理。在调试过程中，我们发现了剪枝条件过严导致路径丢失的问题，通过将">="改为">"并修正启发式计算解决，这让我们深刻理解了算法实现细节的重要性。同时，我们也学会了如何善用AI辅助编程工具（如Copilot）——它能大幅提升代码生成效率，但也要具备审查和调试代码的能力。'),
    ('在工程层面', 
     '我们实践了模块化设计流程，从需求分析、模块划分、代码实现到测试调试，每一步都需要团队协作。我们制定了统一的代码规范，使用extern全局变量实现模块间通信，设计了清晰的接口。在版本控制方面，我们使用Git管理代码，避免了代码冲突。'),
    ('在测试层面', 
     '我们学会了系统性测试的重要性，不仅是正常功能测试，还包括边界条件（起点终点关闭、文件异常、非法输入等）。通过编写自动化测试脚本（Python），我们能够快速回归验证，提升了开发效率。'),
    ('团队收获', 
     '组内每位成员承担了不同模块，有的负责数据加载，有的负责算法，有的负责界面，有的负责测试。通过互相Code Review来完善代码，我们不仅提升了编程水平，也增强了沟通与协作能力。特别是当算法出现Bug时，通过集体讨论和逐步调试，最终找到了根源并修复——这种成就感令人难忘。'),
    ('未来展望', 
     '虽然我们完成了课程设计的基本要求，但仍有不少进阶功能未实现，如换乘站整体关闭、线路停运、网络连通性分析等。如有更多时间，我们希望能够补全这些功能，并尝试使用Qt等框架开发图形界面，使系统更加直观易用。'),
]

for title, content in summary_sections:
    add_para(doc, title + '：', bold=True)
    add_para(doc, content)
    add_para(doc, '')

add_para(doc, '')
add_para(doc, '总的来说，这次课程设计让我们将课堂所学的理论知识真正落地到实际工程问题中，既巩固了数据结构知识，也锻炼了团队合作和项目管理能力，受益匪浅。')

# ===== 保存 =====
output_path = r'c:\Users\MapleLeaf\Desktop\数据结构课设\1\小组报告_规范版.docx'
doc.save(output_path)
print(f'报告已生成：{output_path}')
print('完成！')
